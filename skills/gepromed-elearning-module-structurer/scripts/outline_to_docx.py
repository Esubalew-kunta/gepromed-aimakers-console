#!/usr/bin/env python3
"""Render a GEPROMED e-learning outline to a branded .docx — deterministically.

This is a *pure formatter*. It converts a structured text/markdown outline (the
skill's output format, or simple markdown headings/bullets) into a Word document
with GEPROMED heading styles, a title block, and the bundled logo if present. It
adds **no** content and makes **no** pedagogical decisions — all structure comes
from the input. Use it when the Education team wants an editable course-map doc.

Recognised input lines (leading whitespace allowed):
    # Title / ## / ###             markdown headings  → Heading 1/2/3
    MODULE <...>                    → Heading 1 (blue)
    CHAPTER <...> / CHAPITRE <...>  → Heading 2 (blue)
    - Section ... / * ...           → bullet
    ✓ ...                           → assessment line (orange marker)
    Objective(s): / Objectif(s):    → italic objective line
    everything else                 → normal paragraph

Usage:
    python outline_to_docx.py --in outline.md --out module.docx --title "Vascular e-learning"
    echo "<outline>" | python outline_to_docx.py --out module.docx

Requires: python-docx  (pip install python-docx)
Exit 0 on success, 2 on bad input, 3 if python-docx is missing.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_LOGO = SCRIPT_DIR.parent / "assets" / "gepromed-logo.png"

# GEPROMED brand colours
BLUE = (0x00, 0x7A, 0xC2)
ORANGE = (0xEC, 0x6C, 0x17)
DARK = (0x1F, 0x2A, 0x33)
MUTED = (0x5F, 0x6B, 0x73)


def _rgb(color):
    from docx.shared import RGBColor

    return RGBColor(*color)


def classify(line: str):
    """Return (kind, text) for a stripped line."""
    s = line.strip()
    if not s:
        return ("blank", "")
    if s.startswith("### "):
        return ("h3", s[4:].strip())
    if s.startswith("## "):
        return ("h2", s[3:].strip())
    if s.startswith("# "):
        return ("h1", s[2:].strip())
    if s.startswith("MODULE"):
        return ("h1", s)
    if s.startswith("CHAPTER") or s.startswith("CHAPITRE"):
        return ("h2", s)
    if s.startswith("✓"):
        return ("assess", s.lstrip("✓ ").strip())
    if s.startswith(("Objective", "Objectif")):
        return ("objective", s)
    if s.startswith(("- ", "* ")):
        return ("bullet", s[2:].strip())
    return ("text", s)


def build(lines, out_path: Path, title: str, logo: Path) -> None:
    import docx
    from docx.shared import Pt, Inches

    doc = docx.Document()

    if logo and logo.exists():
        try:
            doc.add_picture(str(logo), width=Inches(1.6))
        except Exception:
            pass  # never fail the doc over a logo

    if title:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.color.rgb = _rgb(BLUE)
        sub = doc.add_paragraph("GEPROMED — René Kieny Education Center · course structure (draft)")
        for run in sub.runs:
            run.font.color.rgb = _rgb(MUTED)
            run.font.size = Pt(9)

    for raw in lines:
        kind, text = classify(raw)
        if kind == "blank":
            continue
        if kind in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[kind]
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.color.rgb = _rgb(BLUE)
        elif kind == "objective":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = _rgb(DARK)
        elif kind == "assess":
            p = doc.add_paragraph()
            marker = p.add_run("✓ ")
            marker.bold = True
            marker.font.color.rgb = _rgb(ORANGE)
            body = p.add_run(text)
            body.font.color.rgb = _rgb(DARK)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)

    doc.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a GEPROMED e-learning outline to .docx.")
    parser.add_argument("--in", dest="infile", default="", help="Input outline file (default: stdin).")
    parser.add_argument("--out", required=True, help="Output .docx path.")
    parser.add_argument("--title", default="GEPROMED e-learning module", help="Document title.")
    parser.add_argument("--logo", default=str(DEFAULT_LOGO), help="Logo image path (optional).")
    args = parser.parse_args()

    if args.infile:
        p = Path(args.infile)
        if not p.exists():
            print(f"ERROR: input file not found: {p}", file=sys.stderr)
            return 2
        lines = p.read_text(encoding="utf-8").splitlines()
    else:
        data = sys.stdin.read()
        if not data.strip():
            print("ERROR: no input on stdin and no --in file given.", file=sys.stderr)
            return 2
        lines = data.splitlines()

    try:
        build(lines, Path(args.out), args.title, Path(args.logo))
    except ImportError:
        print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
        return 3

    print(f"DOCX written → {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
