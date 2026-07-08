#!/usr/bin/env python3
"""Generate a Qualiopi-compliant GEPROMED training PROGRAM as a brand-styled .docx.

Deterministic renderer. The model supplies the *content* (validated against
references/qualiopi-checklist.md); this script lays it out with the GEPROMED
charte: blue headings (#007AC2), the bundled logo, clean typography, and every
RNQ-required block. It does NOT invent content — missing blocks render their
provided value or a bracketed placeholder. A human (RQ) validates before publish.

Input is a JSON file describing the program (see --print-schema). Any required
block left empty is rendered as a bracketed placeholder so the gap is visible.

Usage:
    python generate_program_docx.py --in program.json --out program.docx
    python generate_program_docx.py --print-schema          # show the JSON schema
    python generate_program_docx.py --demo --out demo.docx  # render a sample FR program

Exit 0 on success, 2 on bad input.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).resolve().parent
LOGO = SCRIPT_DIR.parent / "assets" / "gepromed-logo.png"

BLUE = RGBColor(0x00, 0x7A, 0xC2)
ORANGE = RGBColor(0xEC, 0x6C, 0x17)
DARK = RGBColor(0x1F, 0x2A, 0x33)
MUTED = RGBColor(0x5F, 0x6B, 0x73)
TINT = "E1F0F9"

FONT = "Calibri"

# RNQ-required blocks, in render order. (key, FR label, EN label, required)
BLOCKS = [
    ("public_vise", "Public visé", "Target audience", True),
    ("prerequis", "Prérequis", "Prerequisites", True),
    ("objectifs", "Objectifs pédagogiques", "Learning objectives", True),
    ("contenu", "Contenu / programme détaillé", "Detailed content", True),
    ("duree", "Durée", "Duration", True),
    ("modalites_pedagogiques", "Modalités pédagogiques", "Delivery methods", True),
    ("moyens", "Moyens techniques et encadrement", "Resources & trainers", False),
    ("evaluation", "Modalités d'évaluation", "Assessment methods", True),
    ("sanction", "Sanction / validation", "Certificate / proof", False),
    ("accessibilite", "Accessibilité handicap", "Accessibility (disability)", True),
    ("delais_acces", "Délais d'accès", "Access lead time", True),
    ("tarifs", "Tarifs", "Pricing", True),
    ("inscription", "Modalités et délais d'inscription", "Enrolment process", True),
    ("contact", "Contact / référent pédagogique", "Pedagogical contact", False),
    ("indicateurs", "Indicateurs de résultats", "Results indicators", False),
]

DEFAULT_ACCESSIBILITE_FR = (
    "GEPROMED s'engage à étudier toute situation de handicap afin d'envisager "
    "l'adaptation de la formation et des modalités d'accueil. Un référent handicap "
    "est à votre disposition pour analyser votre demande au cas par cas et "
    "déterminer les aménagements possibles. Contact référent handicap : "
    "[email / téléphone à confirmer]."
)
DEFAULT_ACCESSIBILITE_EN = (
    "GEPROMED is committed to reviewing any disability situation in order to adapt "
    "the training and the reception arrangements. A disability referent is available "
    "to assess your request on a case-by-case basis and determine possible "
    "accommodations. Disability referent contact: [email / phone to confirm]."
)
DEFAULT_SANCTION_FR = (
    "Attestation de fin de formation remise à chaque participant. "
    "Certificat de réalisation établi pour les actions concernées."
)
DEFAULT_SANCTION_EN = (
    "Certificate of completion issued to each participant. "
    "Certificate of attendance provided for the relevant actions."
)


def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    for sname in ("Heading 1", "Heading 2", "Title"):
        try:
            doc.styles[sname].font.name = FONT
        except KeyError:
            pass


def _heading(doc: Document, text: str, size: int = 13) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = BLUE
    run.font.name = FONT
    # thin blue rule under the heading
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "007AC2")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _body(doc: Document, text: str, muted: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED if muted else DARK
    p.paragraph_format.space_after = Pt(4)


def _bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(it))
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(2)


def _render_value(doc: Document, value, fr: bool) -> None:
    """Render a block value that may be a string, a list, or a bracket placeholder."""
    if value is None or value == "" or value == []:
        ph = "[À compléter — valider par le RQ]" if fr else "[To complete — RQ to validate]"
        _body(doc, ph, muted=True)
    elif isinstance(value, list):
        _bullets(doc, value)
    else:
        _body(doc, str(value))


def _header_band(doc: Document, data: dict, fr: bool) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO), width=Inches(2.2))
    tagline = doc.add_paragraph()
    r = tagline.add_run("GEPROMED — The medical device hub for patient safety")
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    r.italic = True
    tagline.paragraph_format.space_after = Pt(8)

    # Program title
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(2)
    tr = title.add_run(data.get("intitule") or ("[Intitulé de la formation]" if fr else "[Course title]"))
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = DARK
    tr.font.name = FONT

    sub = doc.add_paragraph()
    label = "Programme de formation" if fr else "Training programme"
    sr = sub.add_run(label)
    sr.font.size = Pt(11)
    sr.font.color.rgb = ORANGE
    sr.bold = True
    sub.paragraph_format.space_after = Pt(2)

    # reference / version / date line
    ref_bits = []
    if data.get("reference"):
        ref_bits.append(f"Réf. {data['reference']}" if fr else f"Ref. {data['reference']}")
    ref_bits.append((f"Version {data.get('version','[v.]')}"))
    ref_bits.append((f"Mise à jour : {data.get('date','[date]')}") if fr else (f"Updated: {data.get('date','[date]')}"))
    meta = doc.add_paragraph()
    mr = meta.add_run("  ·  ".join(ref_bits))
    mr.font.size = Pt(8.5)
    mr.font.color.rgb = MUTED
    meta.paragraph_format.space_after = Pt(8)


def _footer(doc: Document, fr: bool) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    txt = ("GEPROMED — organisme de formation certifié Qualiopi · ISO 9001 · ISO 13485   |   "
           "Document de travail — à valider par le Responsable Qualité avant diffusion") if fr else (
           "GEPROMED — Qualiopi · ISO 9001 · ISO 13485 certified training provider   |   "
           "Working document — to be validated by the Quality Manager before release")
    r = p.add_run(txt)
    r.font.size = Pt(7.5)
    r.font.color.rgb = MUTED


def build(data: dict, out_path: Path) -> Path:
    fr = (data.get("language", "fr").lower() != "en")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    _base_styles(doc)
    _header_band(doc, data, fr)
    _footer(doc, fr)

    # apply GEPROMED defaults for safe blocks if absent
    if not data.get("accessibilite"):
        data["accessibilite"] = DEFAULT_ACCESSIBILITE_FR if fr else DEFAULT_ACCESSIBILITE_EN
    if not data.get("sanction"):
        data["sanction"] = DEFAULT_SANCTION_FR if fr else DEFAULT_SANCTION_EN

    for key, fr_label, en_label, required in BLOCKS:
        value = data.get(key)
        # skip optional blocks that have no value (15/16 etc.) unless required
        if not required and (value is None or value == "" or value == []):
            if key in ("moyens", "sanction", "contact", "indicateurs"):
                continue
        _heading(doc, fr_label if fr else en_label)
        _render_value(doc, value, fr)

    # closing validation note
    doc.add_paragraph()
    note = doc.add_paragraph()
    nt = note.add_run(
        ("Note de conformité : ce programme suit les exigences du Référentiel National "
         "Qualité (Qualiopi). Vérifier que chaque objectif est évaluable et couvert par "
         "les modalités d'évaluation. Toute valeur entre crochets doit être confirmée par "
         "le Responsable Qualité avant diffusion publique.") if fr else
        ("Compliance note: this programme follows the Référentiel National Qualité (Qualiopi) "
         "requirements. Check that each objective is assessable and covered by the assessment "
         "methods. Any bracketed value must be confirmed by the Quality Manager before public "
         "release."))
    nt.font.size = Pt(8.5)
    nt.italic = True
    nt.font.color.rgb = MUTED

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


SCHEMA = {
    "language": "fr | en (default fr)",
    "intitule": "Course title (string)",
    "reference": "internal ref (optional)",
    "version": "version label (optional)",
    "date": "last-updated date (optional)",
    "public_vise": "string or list",
    "prerequis": "string ('Aucun prérequis' if none)",
    "objectifs": "list of operational, assessable objectives (action verbs)",
    "contenu": "list of modules/sessions",
    "duree": "string e.g. '2 jours — 14 heures'",
    "modalites_pedagogiques": "string or list (présentiel/distanciel/simulation/mixte)",
    "moyens": "string or list (optional)",
    "evaluation": "string or list (pre/post-test, grille, satisfaction)",
    "sanction": "string (optional, default attestation)",
    "accessibilite": "string (optional, default GEPROMED process text)",
    "delais_acces": "string",
    "tarifs": "string or list",
    "inscription": "string",
    "contact": "string (optional)",
    "indicateurs": "string or list (optional, only if real figures)",
}

DEMO = {
    "language": "fr",
    "intitule": "Bootcamp Vasculaire — abord et anastomose sur simulateur",
    "reference": "GEP-FORM-VASC-01",
    "version": "1.0",
    "date": "2026-06-20",
    "public_vise": [
        "Chirurgiens vasculaires en exercice",
        "Internes en chirurgie vasculaire (à partir de la 3e année)",
        "Praticiens en formation continue souhaitant consolider leurs gestes d'anastomose",
    ],
    "prerequis": "Statut de praticien ou d'interne en chirurgie vasculaire. Aucun prérequis académique supplémentaire.",
    "objectifs": [
        "Réaliser une anastomose termino-latérale sur simulateur dans le temps imparti.",
        "Identifier et corriger les défauts de suture vasculaire les plus fréquents.",
        "Appliquer les principes d'exposition et de préparation du champ opératoire.",
        "Évaluer la qualité d'une anastomose à l'aide d'une grille standardisée.",
    ],
    "contenu": [
        "Demi-journée 1 — Rappels d'anatomie chirurgicale et principes d'abord (exposé + démonstration).",
        "Demi-journée 2 — Ateliers pratiques de suture vasculaire sur simulateur (dry-lab).",
        "Demi-journée 3 — Anastomoses termino-latérales : mise en situation et débriefing individualisé.",
        "Demi-journée 4 — Évaluation pratique sur grille et synthèse des axes de progression.",
    ],
    "duree": "2 jours — 14 heures (4 demi-journées).",
    "modalites_pedagogiques": [
        "Présentiel au René Kieny Education Center (Strasbourg).",
        "Formation par simulation : ateliers pratiques sur simulateurs vasculaires (dry-lab).",
        "Pédagogie active : démonstrations, mises en situation, débriefing individualisé.",
    ],
    "moyens": [
        "Simulateurs vasculaires et consommables d'entraînement fournis.",
        "Encadrement par des chirurgiens vasculaires formateurs (ratio encadrant/participant adapté à la pratique).",
    ],
    "evaluation": [
        "Pré-test et post-test de connaissances.",
        "Évaluation pratique sur simulateur à l'aide d'une grille d'évaluation standardisée.",
        "Questionnaire de satisfaction à chaud ; questionnaire à froid à [N] semaines.",
    ],
    "duree_note": None,
    "delais_acces": "Inscription possible jusqu'à [N] jours avant la session, dans la limite des places disponibles.",
    "tarifs": "Tarif : [montant] € net de taxe par participant (GEPROMED, organisme à but non lucratif). Prises en charge possibles : OPCO, employeur, financement personnel. Conditions détaillées sur demande.",
    "inscription": "Inscription par formulaire en ligne ou par email auprès du référent pédagogique, jusqu'à [date limite].",
    "contact": "Référent pédagogique GEPROMED — René Kieny Education Center : [email / téléphone à confirmer].",
}


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate a Qualiopi GEPROMED training program .docx")
    ap.add_argument("--in", dest="infile", help="Path to program JSON.")
    ap.add_argument("--out", default="program.docx", help="Output .docx path.")
    ap.add_argument("--print-schema", action="store_true", help="Print the JSON schema and exit.")
    ap.add_argument("--demo", action="store_true", help="Render a bundled FR demo program.")
    args = ap.parse_args()

    if args.print_schema:
        print(json.dumps(SCHEMA, ensure_ascii=False, indent=2))
        return 0

    if args.demo:
        data = DEMO
    elif args.infile:
        try:
            data = json.loads(Path(args.infile).read_text(encoding="utf-8"))
        except Exception as e:  # noqa: BLE001
            print(f"ERROR reading {args.infile}: {e}")
            return 2
    else:
        print("ERROR: provide --in program.json (or --demo / --print-schema).")
        return 2

    out = build(data, Path(args.out))
    print(f"OK — wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
