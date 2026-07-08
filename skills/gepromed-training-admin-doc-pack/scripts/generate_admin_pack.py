#!/usr/bin/env python3
"""Generate a GEPROMED training ADMIN PACK from a participant list + session details.

Produces three brand-styled artifacts in one run:
  1. feuille_emargement.docx — attendance / signing sheet (one row per participant)
  2. badges.docx              — printable name badges (grid on A4); Pillow PNGs optional
  3. notice_rgpd.docx         — the RGPD data-protection notice (FR/EN)

Deterministic renderer following references/layout-spec.md and
references/rgpd-notice-template.md. It NEVER invents personal data, DPO contact,
or retention periods — unknowns render as `[crochets]` for the DPO to confirm.
A human (the DPO for the notice) validates before use.

Input is a JSON file (see --print-schema) or use --demo. Participants is a list of
{nom, prenom, organisme?}.

Usage:
    python generate_admin_pack.py --in session.json --outdir ./pack
    python generate_admin_pack.py --demo --outdir ./pack
    python generate_admin_pack.py --print-schema
    python generate_admin_pack.py --demo --outdir ./pack --badge-png   # also emit PNG badges via Pillow

Exit 0 on success, 2 on bad input.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).resolve().parent
LOGO = SCRIPT_DIR.parent / "assets" / "gepromed-logo.png"

BLUE = RGBColor(0x00, 0x7A, 0xC2)
ORANGE = RGBColor(0xEC, 0x6C, 0x17)
DARK = RGBColor(0x1F, 0x2A, 0x33)
MUTED = RGBColor(0x5F, 0x6B, 0x73)
BLUE_HEX = "007AC2"
FONT = "Calibri"


# ----------------------------- shared helpers -----------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_border(cell, color=BLUE_HEX, sz="6") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def _base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK


def _logo_header(doc: Document, title: str, accent_subtitle: str | None = None) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.add_run().add_picture(str(LOGO), width=Inches(1.9))
    tagline = doc.add_paragraph()
    tr = tagline.add_run("GEPROMED — The medical device hub for patient safety")
    tr.font.size = Pt(8.5)
    tr.italic = True
    tr.font.color.rgb = MUTED
    tagline.paragraph_format.space_after = Pt(6)

    h = doc.add_paragraph()
    hr = h.add_run(title)
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = BLUE
    h.paragraph_format.space_after = Pt(2)
    if accent_subtitle:
        s = doc.add_paragraph()
        sr = s.add_run(accent_subtitle)
        sr.font.size = Pt(10)
        sr.font.color.rgb = ORANGE
        sr.bold = True
        s.paragraph_format.space_after = Pt(6)


def _footer(doc: Document, text: str) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(7.5)
    r.font.color.rgb = MUTED


def _ph(value, placeholder: str) -> str:
    return str(value) if value not in (None, "", []) else placeholder


# ----------------------------- attendance sheet -----------------------------

def build_attendance(data: dict, out: Path, fr: bool) -> Path:
    doc = Document()
    for s in doc.sections:
        s.orientation = 1  # landscape would be 1; keep portrait default
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)
    _base_styles(doc)
    _logo_header(doc,
                 "Feuille d'émargement" if fr else "Attendance sheet",
                 None)

    sess = data.get("session", {})
    info_lines = []
    if fr:
        info_lines.append(("Formation : ", _ph(sess.get("intitule"), "[intitulé de la formation]")))
        info_lines.append(("Date(s) : ", _ph(sess.get("date"), "[date]")))
        info_lines.append(("Horaires : ", _ph(sess.get("horaires"), "[horaires]")))
        info_lines.append(("Lieu : ", _ph(sess.get("lieu"), "[lieu]")))
        info_lines.append(("Formateur : ", _ph(sess.get("formateur"), "[formateur]")))
    else:
        info_lines.append(("Course: ", _ph(sess.get("intitule"), "[course title]")))
        info_lines.append(("Date(s): ", _ph(sess.get("date"), "[date]")))
        info_lines.append(("Hours: ", _ph(sess.get("horaires"), "[hours]")))
        info_lines.append(("Venue: ", _ph(sess.get("lieu"), "[venue]")))
        info_lines.append(("Trainer: ", _ph(sess.get("formateur"), "[trainer]")))
    for label, val in info_lines:
        p = doc.add_paragraph()
        lr = p.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = DARK
        vr = p.add_run(val)
        vr.font.size = Pt(10)
        vr.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if fr:
        headers = ["#", "Nom", "Prénom", "Organisme / Fonction", "Signature matin", "Signature après-midi"]
    else:
        headers = ["#", "Last name", "First name", "Organisation / Role", "Signature AM", "Signature PM"]

    participants = data.get("participants", [])
    n = max(len(participants), 1)
    table = doc.add_table(rows=1 + n, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for j, htxt in enumerate(headers):
        hdr[j].text = ""
        _set_cell_bg(hdr[j], BLUE_HEX)
        _cell_border(hdr[j])
        run = hdr[j].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(n):
        row = table.rows[i + 1].cells
        p = participants[i] if i < len(participants) else {}
        values = [
            str(i + 1),
            p.get("nom", ""),
            p.get("prenom", ""),
            p.get("organisme", ""),
            "",  # signature AM
            "",  # signature PM
        ]
        for j, v in enumerate(values):
            cell = row[j]
            cell.text = ""
            _cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].add_run(v)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
            if j == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # generous row height for the signature
        trPr = table.rows[i + 1]._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "520")
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig_txt = "Fait à __________________ , le __________________     Signature du formateur :" if fr \
        else "Done at __________________ , on __________________     Trainer's signature:"
    sr = sig.add_run(sig_txt)
    sr.font.size = Pt(9.5)
    sr.font.color.rgb = DARK

    _footer(doc, ("Document de présence — GEPROMED, organisme de formation certifié Qualiopi · "
                  "À conserver pour le suivi qualité.") if fr else
                 ("Attendance record — GEPROMED, Qualiopi-certified training provider · "
                  "Keep for quality follow-up."))
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ----------------------------- badges (docx grid) -----------------------------

def build_badges_docx(data: dict, out: Path, fr: bool) -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.0)
        s.bottom_margin = Cm(1.0)
        s.left_margin = Cm(1.0)
        s.right_margin = Cm(1.0)
    _base_styles(doc)

    participants = data.get("participants", [])
    intitule = data.get("session", {}).get("intitule", "")
    cols = 2
    rows = (len(participants) + cols - 1) // cols if participants else 1
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    idx = 0
    for r in range(rows):
        # fix row height ~ badge height
        trPr = table.rows[r]._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "1500")
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)
        for c in range(cols):
            cell = table.rows[r].cells[c]
            cell.text = ""
            _cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx < len(participants):
                p = participants[idx]
                # logo small
                if LOGO.exists():
                    lp = cell.paragraphs[0]
                    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    lp.add_run().add_picture(str(LOGO), width=Inches(1.0))
                # name
                np = cell.add_paragraph()
                np.alignment = WD_ALIGN_PARAGRAPH.CENTER
                name = f"{p.get('prenom','').strip()} {p.get('nom','').strip().upper()}".strip()
                nr = np.add_run(name or "[Nom]")
                nr.bold = True
                nr.font.size = Pt(16)
                nr.font.color.rgb = DARK
                # org
                if p.get("organisme"):
                    op = cell.add_paragraph()
                    op.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    orr = op.add_run(str(p["organisme"]))
                    orr.font.size = Pt(10)
                    orr.font.color.rgb = MUTED
                # course line
                if intitule:
                    cp = cell.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(intitule)
                    cr.font.size = Pt(8)
                    cr.font.color.rgb = BLUE
            idx += 1
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def build_badges_png(data: dict, outdir: Path) -> list[Path]:
    """Optional crisp PNG badges via Pillow."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:  # noqa: BLE001
        print("  (Pillow not available — skipping PNG badges)")
        return []
    participants = data.get("participants", [])
    intitule = data.get("session", {}).get("intitule", "")
    W, H = 1063, 650  # ~90x55mm at 300dpi
    made = []

    def _font(size):
        for cand in ("DejaVuSans-Bold.ttf", "DejaVuSans.ttf"):
            try:
                return ImageFont.truetype(cand, size)
            except Exception:  # noqa: BLE001
                continue
        return ImageFont.load_default()

    logo_img = None
    if LOGO.exists():
        try:
            logo_img = Image.open(LOGO).convert("RGBA")
        except Exception:  # noqa: BLE001
            logo_img = None

    badge_dir = outdir / "badges_png"
    badge_dir.mkdir(parents=True, exist_ok=True)
    for i, p in enumerate(participants):
        img = Image.new("RGB", (W, H), "white")
        d = ImageDraw.Draw(img)
        d.rectangle([6, 6, W - 6, H - 6], outline=(0, 122, 194), width=6)
        y = 40
        if logo_img is not None:
            lw = 360
            ratio = lw / logo_img.width
            lh = int(logo_img.height * ratio)
            lg = logo_img.resize((lw, lh))
            img.paste(lg, ((W - lw) // 2, y), lg)
            y += lh + 30
        name = f"{p.get('prenom','').strip()} {p.get('nom','').strip().upper()}".strip() or "[Nom]"
        f_name = _font(72)
        tb = d.textbbox((0, 0), name, font=f_name)
        d.text(((W - (tb[2] - tb[0])) // 2, y), name, fill=(31, 42, 51), font=f_name)
        y += (tb[3] - tb[1]) + 30
        if p.get("organisme"):
            f_org = _font(40)
            org = str(p["organisme"])
            ob = d.textbbox((0, 0), org, font=f_org)
            d.text(((W - (ob[2] - ob[0])) // 2, y), org, fill=(95, 107, 115), font=f_org)
        if intitule:
            f_c = _font(30)
            cb = d.textbbox((0, 0), intitule, font=f_c)
            d.text(((W - (cb[2] - cb[0])) // 2, H - 70), intitule, fill=(0, 122, 194), font=f_c)
        out = badge_dir / f"badge_{i+1:02d}.png"
        img.save(out)
        made.append(out)
    return made


# ----------------------------- RGPD notice -----------------------------

NOTICE_FR = [
    ("Dans le cadre de votre inscription et de votre participation à la formation "
     "« {intitule} » organisée par GEPROMED, vos données personnelles (nom, prénom, "
     "coordonnées, profession/établissement, émargement) font l'objet d'un traitement.", None),
    ("Responsable de traitement : ", "GEPROMED — {adresse}, Strasbourg."),
    ("Délégué à la protection des données (DPO) : ", "{dpo}."),
    ("Finalités : ", "gestion administrative et pédagogique de la formation, suivi de "
     "l'assiduité (émargement), délivrance des attestations, et obligations liées à la "
     "certification Qualiopi."),
    ("Base légale : ", "exécution de mesures précontractuelles et contractuelles, et respect "
     "d'obligations légales et réglementaires de l'organisme de formation."),
    ("Destinataires : ", "les services internes habilités de GEPROMED et, le cas échéant, les "
     "financeurs (OPCO, employeur) et les autorités de contrôle."),
    ("Durée de conservation : ", "{retention}."),
    ("Vos droits : ", "vous disposez d'un droit d'accès, de rectification, d'effacement, de "
     "limitation, d'opposition et de portabilité de vos données. Vous pouvez les exercer en "
     "écrivant à {dpo}. Vous avez également le droit d'introduire une réclamation auprès de la "
     "CNIL (www.cnil.fr)."),
    ("Photographies / vidéos : ", "toute prise d'image fait l'objet d'un consentement distinct ; "
     "vous pouvez le refuser sans conséquence sur votre participation."),
]

NOTICE_EN = [
    ("As part of your registration and participation in the training course \"{intitule}\" "
     "organised by GEPROMED, your personal data (last name, first name, contact details, "
     "profession/organisation, attendance signature) is processed.", None),
    ("Data controller: ", "GEPROMED — {adresse}, Strasbourg."),
    ("Data Protection Officer (DPO): ", "{dpo}."),
    ("Purposes: ", "administrative and pedagogical management of the training, attendance "
     "tracking, issuance of certificates, and obligations related to the Qualiopi certification."),
    ("Legal basis: ", "performance of pre-contractual and contractual measures, and compliance "
     "with the legal and regulatory obligations of the training provider."),
    ("Recipients: ", "authorised internal GEPROMED staff and, where applicable, funders (OPCO, "
     "employer) and supervisory authorities."),
    ("Retention period: ", "{retention}."),
    ("Your rights: ", "you have the right to access, rectify, erase, restrict, object to, and "
     "port your data. You may exercise these rights by writing to {dpo}. You also have the right "
     "to lodge a complaint with the CNIL (www.cnil.fr)."),
    ("Photographs / video: ", "any image capture is subject to separate consent; you may decline "
     "without affecting your participation."),
]


def build_rgpd(data: dict, out: Path, fr: bool) -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    _base_styles(doc)
    _logo_header(doc,
                 "Information sur le traitement de vos données personnelles" if fr else
                 "Information on the processing of your personal data")

    sess = data.get("session", {})
    rgpd = data.get("rgpd", {})
    fields = {
        "intitule": _ph(sess.get("intitule"), "[intitulé de la formation]" if fr else "[course title]"),
        "adresse": _ph(rgpd.get("adresse"), "[adresse]" if fr else "[address]"),
        "dpo": _ph(rgpd.get("dpo_email"), "[email DPO à confirmer]" if fr else "[DPO email to confirm]"),
        "retention": _ph(rgpd.get("retention"),
                         "[durée à confirmer par le DPO]" if fr else "[period to be confirmed by the DPO]"),
    }
    blocks = NOTICE_FR if fr else NOTICE_EN
    for label, body in blocks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if body is None:
            run = p.add_run(label.format(**fields))
            run.font.size = Pt(10.5)
            run.font.color.rgb = DARK
        else:
            lr = p.add_run(label)
            lr.bold = True
            lr.font.size = Pt(10.5)
            lr.font.color.rgb = BLUE
            br = p.add_run(body.format(**fields))
            br.font.size = Pt(10.5)
            br.font.color.rgb = DARK

    _footer(doc, "À valider par le DPO avant diffusion." if fr else
                 "To be validated by the DPO before release.")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ----------------------------- driver -----------------------------

SCHEMA = {
    "language": "fr | en (default fr)",
    "session": {
        "intitule": "course title",
        "date": "date(s)",
        "horaires": "hours e.g. 9h-17h",
        "lieu": "venue",
        "formateur": "trainer name/role",
    },
    "participants": [{"nom": "LAST", "prenom": "First", "organisme": "org/role (optional)"}],
    "rgpd": {
        "adresse": "GEPROMED address (optional)",
        "dpo_email": "DPO contact (optional — bracketed if absent)",
        "retention": "retention period (optional — bracketed if absent)",
    },
}

DEMO = {
    "language": "fr",
    "session": {
        "intitule": "Bootcamp Vasculaire — anastomose sur simulateur",
        "date": "12-13 octobre 2026",
        "horaires": "9h00 - 17h00",
        "lieu": "René Kieny Education Center, Strasbourg",
        "formateur": "Dr [Nom], chirurgien vasculaire formateur",
    },
    "participants": [
        {"nom": "MARTIN", "prenom": "Claire", "organisme": "CHU de Strasbourg"},
        {"nom": "DUBOIS", "prenom": "Hugo", "organisme": "Interne — CHU de Lyon"},
        {"nom": "BERNARD", "prenom": "Aïcha", "organisme": "Clinique Pasteur"},
        {"nom": "PETIT", "prenom": "Lucas", "organisme": ""},
    ],
    "rgpd": {
        "adresse": "[adresse]",
        "dpo_email": "[email DPO à confirmer]",
        "retention": "[durée à confirmer par le DPO]",
    },
}


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate a GEPROMED training admin pack.")
    ap.add_argument("--in", dest="infile", help="Path to session JSON.")
    ap.add_argument("--outdir", default="./pack", help="Output directory.")
    ap.add_argument("--print-schema", action="store_true")
    ap.add_argument("--demo", action="store_true")
    ap.add_argument("--badge-png", action="store_true", help="Also emit PNG badges via Pillow.")
    args = ap.parse_args()

    if args.print_schema:
        print(json.dumps(SCHEMA, ensure_ascii=False, indent=2))
        return 0

    if args.demo:
        data = DEMO
    elif args.infile:
        try:
            data = json.loads(Path(args.infile).read_text(encoding="utf-8"))
        except Exception as e:  # noqa: BLE001
            print(f"ERROR reading {args.infile}: {e}")
            return 2
    else:
        print("ERROR: provide --in session.json (or --demo / --print-schema).")
        return 2

    fr = (data.get("language", "fr").lower() != "en")
    outdir = Path(args.outdir)
    a = build_attendance(data, outdir / "feuille_emargement.docx", fr)
    b = build_badges_docx(data, outdir / "badges.docx", fr)
    c = build_rgpd(data, outdir / "notice_rgpd.docx", fr)
    print(f"OK — wrote:\n  {a}\n  {b}\n  {c}")
    if args.badge_png:
        pngs = build_badges_png(data, outdir)
        for p in pngs:
            print(f"  {p}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
